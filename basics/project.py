import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import docx
import time

def create_df_stats_summary(dataframe, features_to_include):
    sub_df = dataframe[features_to_include].copy()
    return sub_df.describe()


def create_scatterplot(dataframe, xaxis, yaxis, colour, plot_name,
                       xaxis_scale=None, yaxis_scale=None):
    fig, ax = plt.subplots()

    ax.scatter(dataframe[xaxis], dataframe[yaxis],
               c=dataframe[colour], cmap='viridis')

    ax.set_xlabel(xaxis)
    ax.set_ylabel(yaxis)

    if xaxis_scale is not None:
        ax.set_xlim(xmin=xaxis_scale[0], xmax=xaxis_scale[1])

    if yaxis_scale is not None:
        ax.set_ylim(ymin=yaxis_scale[0], ymax=yaxis_scale[1])

    filename = f'{plot_name}.png'
    plt.savefig(filename)


def add_df_to_docx(doc, dataframe):
    # Reset the index and get the new shape
    dataframe = dataframe.reset_index()
    num_rows, num_cols = dataframe.shape

    # Add a table to the document with the necessary number
    # of rows and columns
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)

    # Add the header row
    for i, col in enumerate(dataframe.columns):
        table.cell(0, i).text = str(col)

    # Add the data rows
    for i, row in dataframe.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)

    return table


def generate_report(report_title, report_author, report_date, report_client,
                    section_title=None,
                    section_text_summary=None,
                    data_stats_summary=None,
                    graph_figure=None):
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading(report_title, 0)
    doc.add_paragraph(f'Authored By: {report_author}')
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Created For: {report_client}')
    doc.add_heading(section_title, 1)
    doc.add_paragraph(section_text_summary)

    # Add Scatter plot
    doc.add_heading('Data Visualisation', 2)
    doc.add_picture(graph_figure)

    # Add dataframe summary
    doc.add_heading('Data Summary', 2)
    summary_table = add_df_to_docx(doc, data_stats_summary)
    summary_table.style = 'LightShading-Accent1'

    doc.save('report.docx')

    return st.info('Report Generated')


# st.set_page_config(layout='wide')

st.title('Streamlit Data Report Generator')

st.markdown('Use this Streamlit app to make your own scatterplot about Xeek and FORCE 2020!')

df = pd.read_csv('https://raw.githubusercontent.com/andymcdgeo/Petrophysics-Python-Series/master/'
                 'Data/Xeek_Well_15-9-15.csv')

with st.form('report'):
    st.write("### Report Details")
    col1, col2 = st.columns(2, gap='large')

    report_title = col1.text_input("Enter report title")
    report_author = col1.text_input("Enter the report author's name")
    report_date = col2.date_input("Select a date for the report")
    report_client = col2.text_input("Enter the client's name")

    sect_col1, sect_col2 = st.columns(2, gap='large')

    sect_col1.write("### Section Details")
    section_title = sect_col1.text_input("Enter section title")
    section_text_summary = sect_col1.text_area("Section Summary")

    data_features = df.columns

    sect_col2.write("### Data Summary")
    data_to_summarise = sect_col2.multiselect("Select features to include in statistical summary",
                                              options=data_features)

    st.write("### Scatterplot Setup")
    sub_col1, sub_col2, sub_col3 = st.columns(3)

    chart_x = sub_col1.selectbox('X axis', options=data_features)
    chart_y = sub_col2.selectbox('Y axis', options=data_features)
    chart_z = sub_col3.selectbox('Z axis', options=data_features)

    if st.form_submit_button('Generate'):
        summary_stats = create_df_stats_summary(df, data_to_summarise)
        scatter_plot_file = create_scatterplot(df, chart_x, chart_y, chart_z,
                                               plot_name='scatter', yaxis_scale=[3, 1], )

        generate_report(report_title, report_author, report_date, report_client,
                        section_title, section_text_summary, summary_stats,
                        graph_figure='scatter.png')